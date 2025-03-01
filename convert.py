import random
from docx import Document
import pypandoc
import re
import os

UNICODES = {
    "a": ["а"],
    "c": ["с"],
    # "d": ["ԁ"],
    "e": ["е"],
    "i": ["і"],
    "j": ["ј"],
    "n": ["n"],
    "o": ["о", "ο", "օ"],
    "p": ["р"],
    # "u": ["υ", "ս"],
    "v": ["ѵ", "ν"],
    "x": ["x"],
    "y": ["у"],
}


def process_runs(paragraph):
    """Process each run while preserving formatting"""
    for run in paragraph.runs:
        text = run.text
        new_text = replace_with_unicode_lookalikes(text)
        if text != new_text:
            print(f"Changed: '{text}' -> '{new_text}'")
        run.text = new_text


def process_document(input_path, output_path):
    try:
        if os.path.exists(output_path):
            os.remove(output_path)
            print(f"Deleted existing file: {output_path}")

        print("Starting document processing...")
        doc = Document(input_path)

        for paragraph in doc.paragraphs:
            process_runs(paragraph)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_runs(paragraph)
        doc.save(output_path)
        print(f"\nSuccessfully processed document. Output saved to: {output_path}")

    except Exception as e:
        print(f"An error occurred: {str(e)}")


def replace_with_unicode_lookalikes(text):
    result = ""
    was_changed = False
    for char in text:
        lower_char = char.lower()
        if lower_char in UNICODES:
            replacement = random.choice(UNICODES[lower_char])
            result += replacement if char.islower() else replacement.upper()
            was_changed = True
        else:
            result += char
    return result


if __name__ == "__main__":
    input_file = "input.docx"
    output_file = "output.docx"
    process_document(input_file, output_file)
